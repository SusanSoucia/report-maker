#!/usr/bin/env python3
"""
填充工厂方法模式实验报告
基于模板: 1.软件架构与设计模式实验（创建型模式）.docx
"""
import copy
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

TEMPLATE = "/home/susan/MySpace/codeSpace/设计模式实验/raw/1.软件架构与设计模式实验（创建型模式）.docx"
OUTPUT = "/home/susan/MySpace/codeSpace/设计模式实验/output/工厂方法模式实验报告.docx"


def clear_paragraph(para):
    """清空段落中所有 run"""
    for run in para.runs:
        run.text = ""


def set_paragraph_text(para, text, font_name="宋体", font_size=Pt(12), bold=False):
    """清空段落并写入新文本（保留第一个 run 的格式作为基准）"""
    # 清空所有 runs
    for run in para.runs:
        run.text = ""
    # 如果有 run，用第一个；否则添加一个
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run("")
    run.text = text
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold


def add_content_paragraphs(cell, after_idx, lines, font_name="宋体", font_size=Pt(12)):
    """在指定索引后插入新段落，返回最后插入的段落索引"""
    # Get the paragraph element to insert after
    base_para = cell.paragraphs[after_idx]
    last_added = after_idx

    for line in lines:
        # Add a new paragraph by cloning the XML element
        new_para_elem = copy.deepcopy(base_para._element)
        # Clear all runs
        for r in new_para_elem.findall(qn('w:r')):
            new_para_elem.remove(r)
        # Insert after base_para
        base_para._element.addnext(new_para_elem)
        # Create a new Paragraph object
        new_para = docx.text.paragraph.Paragraph(new_para_elem, cell)
        # Add run with text
        run = new_para.add_run(line)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = font_size
        last_added += 1

    return last_added


def main():
    doc = docx.Document(TEMPLATE)

    # ========== Table 1, Row 3: Main content cell ==========
    table1 = doc.tables[1]
    cell = table1.rows[3].cells[0]

    paras = cell.paragraphs
    # P0: 一、实验任务 (keep)
    # P1: task description (keep, it's good)
    # P2: 二、实验环境 (keep title)
    # P3: environment placeholder → replace with actual env
    # P4: empty (keep)
    # P5: 三、实验内容 (keep title)
    # P6: 1. 方法 → replace
    # P7: 2. 步骤 → replace
    # P8: 3. 效果 → replace

    # --- 二、实验环境 (P3) ---
    env_lines = [
        "操作系统：Windows 11 + WSL2 (Ubuntu 24.04, Linux 6.6.87.2-microsoft-standard-WSL2)",
        "编程语言：Java (OpenJDK 21.0.11)",
        "开发环境：Visual Studio Code",
        "建模工具：PlantUML 1.2020.02 + Graphviz 2.43.0",
        "AI 辅助工具：Claude Code (DeepSeek-v4-pro)，用于需求分析、代码生成、UML 建模和报告撰写",
    ]
    env_text = "\n".join(env_lines)

    # --- 三.1 方法 (P6) ---
    method_lines = [
        "本实验采用 AI 辅助实验报告写作的方法。传统实验流程中，学生需手工编写代码、手动绘制 UML 图、自行组织报告结构，各环节耗时且容易出错。本实验探索将 AI 大模型（Claude Code）作为协同工具嵌入实验全流程：AI 负责需求梳理、Java 代码模板生成、PlantUML 类图源码编写和报告初稿填充；学生负责需求确认与修正、代码审查与运行验证、截图记录和最终报告审核。",
        "",
        "核心思路是将 AI 定位为“实验助手”而非“替代者”。具体体现为三个原则：",
        "（1）人机协同确认：AI 产出需求清单后，学生逐项核对并补充个人理解，确保方向一致后再进入代码生成阶段。",
        "（2）产物可复现：所有 AI 生成的中间产物——PlantUML 源码（.puml）、Java 源码（.java）——均以文本文件保存于项目目录。修改设计只需修改源码重新运行，不依赖 AI 重新生成。",
        "（3）设计分析反向验证：AI 根据实际代码实现反推生成设计分析文字，而非凭空编造，确保报告中“为什么这样设计”的论述与代码结构严格对应。",
        "",
        "创新点：（1）将 AI 大模型引入本科实验教学流程，探索人机协同的实验报告写作模式；（2）所有设计产物代码化、版本可控，避免 AI “黑箱”输出的不可信问题；（3）报告中的 UML 类图、设计分析、代码实现三者保持一致性——因为它们共享同一套结构描述。",
    ]
    method_text = "\n".join(method_lines)

    # --- 三.2 步骤 (P7) ---
    steps_lines = [
        "（1）需求确认阶段：将实验要求文档（个人要求.txt）和实验模板提供给 AI。AI 提取模板的章节结构（封面→一、实验任务→二、实验环境→三、实验内容→四、实验步骤）和需求清单（7 个 Java 类/接口 + 1 个 UML 类图 + 运行截图），经学生逐项确认后进入实现阶段。",
        "    [此处插入截图：AI 需求确认对话]",
        "",
        "（2）代码生成阶段：AI 根据确认后的类结构生成 7 个 Java 源文件——TV.java（抽象产品接口）、HaierTV.java / HisenseTV.java（具体产品类）、TVFactory.java（抽象工厂接口）、HaierTVFactory.java / HisenseTVFactory.java（具体工厂类）、Client.java（客户端）。所有文件保持一致的注释风格和命名规范。学生审查代码逻辑后，在终端执行 javac 编译和 java 运行。",
        "    [此处插入截图：src/ 目录下 Java 源文件列表]",
        "",
        "（3）UML 建模阶段：AI 根据代码的类结构和依赖关系生成 PlantUML 类图源码（class_diagram.puml）。类图清晰展示了工厂方法模式的四个核心角色及其关系：TV（抽象产品）、HaierTV / HisenseTV（具体产品）、TVFactory（抽象工厂）、HaierTVFactory / HisenseTVFactory（具体工厂），以及 Client 对抽象工厂和抽象产品的依赖。使用 plantuml 命令渲染为 PNG 图片。",
        "    [此处插入截图：UML 类图 — 见 image/class_diagram.png]",
        "",
        "（4）运行验证阶段：编译执行 Java 程序，输出结果显示海尔工厂成功创建海尔电视机并调用 play()，海信工厂同理。验证了工厂方法模式的核心机制——客户端通过抽象工厂接口获得产品实例，不感知具体产品类的存在。",
        "    [此处插入截图：终端运行结果]",
        "",
        "（5）报告生成阶段：AI 按照模板的章节结构逐章填充实验报告，包括实验任务描述、实验环境、方法思路、步骤过程和效果分析。学生在 Word 中完成最终排版调整、插入截图、更新目录和页码。",
    ]
    steps_text = "\n".join(steps_lines)

    # --- 三.3 效果 (P8) ---
    results_lines = [
        "（1）运行结果：程序成功通过海尔工厂创建海尔电视机、通过海信工厂创建海信电视机，并分别调用 play() 方法输出播放信息。运行结果与预期完全一致。",
        "    [此处插入截图：终端完整运行结果]",
        "",
        "（2）UML 类图：PlantUML 生成的类图完整呈现了工厂方法模式的结构——TV 接口作为抽象产品，HaierTV 和 HisenseTV 实现该接口；TVFactory 接口作为抽象工厂，HaierTVFactory 和 HisenseTVFactory 实现该接口并分别创建对应产品。Client 仅依赖 TVFactory 和 TV 两个抽象，符合依赖倒置原则。",
        "    [此处插入截图：UML 类图（image/class_diagram.png）]",
        "",
        "（3）设计特性验证：",
        "    · 创建与使用分离：Client 通过 TVFactory.createTV() 获得 TV 对象，不直接 new 具体产品，降低耦合。",
        "    · 开闭原则：如需新增 TCL 品牌，只需新增 TCLTV 和 TCLTVFactory，无需修改已有类。模板任务描述中明确提到了这一可扩展性要求，本设计完全满足。",
        "    · 单一职责：每个具体工厂（HaierTVFactory / HisenseTVFactory）只负责创建一种产品，职责边界清晰。",
        "",
        "（4）AI 辅助效果评述：AI 在本次实验中主要贡献在于将重复性工作（代码模板编写、PlantUML 源码生成、报告各章节初稿）自动化，使学生可以将精力集中在理解设计模式原理、审查代码正确性和完善设计分析上。所有 AI 产出均经过人工确认和修改，不是简单的“复制粘贴”。代码和 PlantUML 源码保存在项目目录中，可随时修改和重新生成，具有良好的可复现性。",
    ]
    results_text = "\n".join(results_lines)

    # Apply changes to paragraphs
    set_paragraph_text(paras[3], env_text, font_name="宋体", font_size=Pt(12))
    set_paragraph_text(paras[6], method_text, font_name="宋体", font_size=Pt(12))
    set_paragraph_text(paras[7], steps_text, font_name="宋体", font_size=Pt(12))
    set_paragraph_text(paras[8], results_text, font_name="宋体", font_size=Pt(12))

    # ========== Table 2: 四、实验步骤 ==========
    table2 = doc.tables[2]
    cell2 = table2.rows[0].cells[0]
    paras2 = cell2.paragraphs
    # P0: title (keep)
    # P1: empty → fill
    # P2: empty → fill

    step_lines = [
        "（1）定义 TV 接口（抽象产品）：声明 play():void 方法，作为所有电视机的统一行为契约。",
        "（2）实现 HaierTV 和 HisenseTV 具体产品类：分别实现 TV 接口，在 play() 方法中输出各品牌电视机的播放信息。",
        "（3）定义 TVFactory 接口（抽象工厂）：声明 createTV():TV 工厂方法，返回抽象产品类型。",
        "（4）实现 HaierTVFactory 和 HisenseTVFactory 具体工厂类：分别实现 createTV() 方法，返回对应的具体产品实例（new HaierTV() / new HisenseTV()）。",
        "（5）编写 Client 客户端类：在 main() 方法中，通过 TVFactory 接口变量引用具体工厂实例，调用 createTV() 获得 TV 对象，调用 play() 验证。Client 仅依赖抽象（TVFactory、TV），不依赖具体类。",
        "（6）编译运行：在终端执行 javac *.java 编译所有源文件，执行 java Client 运行程序，观察输出结果。",
        "（7）UML 建模：使用 PlantUML 编写类图源码（class_diagram.puml），描述接口、类及其继承/实现/依赖关系，执行 plantuml -tpng 生成 PNG 图片。",
        "（8）撰写实验报告：按照实验模板结构，整理需求分析、设计思路、实现步骤和运行效果，填入各章节。",
    ]
    steps_text_4 = "\n".join(step_lines)

    # Fill paragraph 1 and 2 in Table 2
    set_paragraph_text(paras2[1], steps_text_4, font_name="宋体", font_size=Pt(12))
    # P2 remains empty or can be used for additional notes

    # Save
    doc.save(OUTPUT)
    print(f"报告已生成: {OUTPUT}")


if __name__ == "__main__":
    main()
